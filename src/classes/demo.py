import hashlib
import datetime
import re
from flask import make_response
from docx import Document


# CLASSES
class Demo:
    def __init__(self,token=None):
        self.token = token
        pass

    def create_demo_token(self):
        value = 'replaceup-21041901-'
        today = datetime.date.today().strftime('%Y%m%d%H')
        token_input = value + today
        self.token = sha256(token_input)

    def receive_multiple_inputs(self,data):
        r_elements = {}
        w_elements = {}
        for key, value in data:
            r_match = re.match(r'r_variable(\d{3})', key)
            w_match = re.match(r'w_variable(\d{3})', key)
            if r_match:
                r_elements[r_match.group(1)] = value
            elif w_match:
                w_elements[w_match.group(1)] = value
        output = {r_elements[key]: w_elements.get(key, '') for key in r_elements}
        return output 


# FUNCTIONS
def sha256(value):
    # Encripta el valor en SHA-256
    sha256_hash = hashlib.sha256(value.encode()).hexdigest()
    return sha256_hash

def multiple_replace(file, replacements):
    doc = Document(file)
    print(replacements)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for search, replace in replacements.items():
                run.text = run.text.replace(search, replace)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for search, replace in replacements.items():
                    cell.text = cell.text.replace(search, replace)
    return doc

def replace(replacements,file_name):
    file_path = 'static/import/'
    temp_name = 'output.docx'
    export_path = 'static/export/' + temp_name
    template = file_path + file_name

    file = multiple_replace(template,replacements)
    file.save(export_path)
    with open(export_path, 'rb') as file:
        file_contents = file.read()
    response = make_response(file_contents)
    response.headers['Content-Disposition'] = f'attachment; filename={file_name}'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    return response